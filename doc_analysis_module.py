#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
文档智能分析模块 - 支持PDF、Word、Excel、图片等多种格式
"""

import os
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Optional

DOC_ANALYSIS_DIR = "doc_analysis_data"
DOC_INDEX_FILE = os.path.join(DOC_ANALYSIS_DIR, "doc_index.json")


def ensure_doc_dir():
    """确保文档分析目录存在"""
    os.makedirs(DOC_ANALYSIS_DIR, exist_ok=True)


def load_json(filepath: str) -> dict:
    """加载JSON文件"""
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    return {}


def save_json(filepath: str, data: dict):
    """保存JSON文件"""
    ensure_doc_dir()
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# ========== 文档解析 ==========

def parse_txt(file_path: str) -> dict:
    """解析TXT文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            "success": True,
            "content": content,
            "word_count": len(content),
            "line_count": content.count('\n') + 1
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_pdf(file_path: str) -> dict:
    """解析PDF文件"""
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            content = ""
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
        
        return {
            "success": True,
            "content": content,
            "page_count": page_count,
            "word_count": len(content)
        }
    except ImportError:
        return {"success": False, "error": "需要安装pdfplumber: pip install pdfplumber"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_docx(file_path: str) -> dict:
    """解析Word文件"""
    try:
        from docx import Document
        doc = Document(file_path)
        content = "\n".join([para.text for para in doc.paragraphs if para.text])
        
        return {
            "success": True,
            "content": content,
            "paragraph_count": len(doc.paragraphs),
            "word_count": len(content)
        }
    except ImportError:
        return {"success": False, "error": "需要安装python-docx: pip install python-docx"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_excel(file_path: str) -> dict:
    """解析Excel文件"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_path, data_only=True)
        
        sheets_data = {}
        total_rows = 0
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows = []
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    rows.append([str(cell) if cell is not None else "" for cell in row])
            sheets_data[sheet_name] = rows
            total_rows += len(rows)
        
        # 生成文本摘要
        content = ""
        for sheet_name, rows in sheets_data.items():
            content += f"\n【{sheet_name}】\n"
            for row in rows[:10]:  # 只显示前10行
                content += " | ".join(row) + "\n"
            if len(rows) > 10:
                content += f"... 共{len(rows)}行\n"
        
        return {
            "success": True,
            "content": content,
            "sheet_count": len(sheets_data),
            "total_rows": total_rows,
            "sheets": list(sheets_data.keys())
        }
    except ImportError:
        return {"success": False, "error": "需要安装openpyxl: pip install openpyxl"}
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_csv(file_path: str) -> dict:
    """解析CSV文件"""
    try:
        import csv
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f)
            rows = list(reader)
        
        content = "\n".join([" | ".join(row) for row in rows[:20]])
        if len(rows) > 20:
            content += f"\n... 共{len(rows)}行"
        
        return {
            "success": True,
            "content": content,
            "row_count": len(rows),
            "column_count": len(rows[0]) if rows else 0
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_json_file(file_path: str) -> dict:
    """解析JSON文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        content = json.dumps(data, ensure_ascii=False, indent=2)
        
        return {
            "success": True,
            "content": content,
            "key_count": len(data) if isinstance(data, dict) else len(data) if isinstance(data, list) else 0
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def parse_markdown(file_path: str) -> dict:
    """解析Markdown文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # 提取标题
        headers = [line.strip('# ') for line in content.split('\n') if line.startswith('#')]
        
        return {
            "success": True,
            "content": content,
            "word_count": len(content),
            "headers": headers[:10]
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


# ========== 文档智能分析 ==========

def analyze_document(file_path: str) -> dict:
    """智能分析文档"""
    if not os.path.exists(file_path):
        return {"success": False, "error": f"文件不存在: {file_path}"}
    
    file_ext = os.path.splitext(file_path)[1].lower()
    file_name = os.path.basename(file_path)
    file_size = os.path.getsize(file_path)
    
    # 根据文件类型选择解析器
    parsers = {
        '.txt': parse_txt,
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.doc': parse_docx,
        '.xlsx': parse_excel,
        '.xls': parse_excel,
        '.csv': parse_csv,
        '.json': parse_json_file,
        '.md': parse_markdown,
    }
    
    parser = parsers.get(file_ext)
    if not parser:
        return {
            "success": False,
            "error": f"不支持的文件格式: {file_ext}",
            "supported_formats": list(parsers.keys())
        }
    
    # 解析文档
    result = parser(file_path)
    
    # 添加文件信息
    result["file_name"] = file_name
    result["file_path"] = file_path
    result["file_size"] = file_size
    result["file_type"] = file_ext
    result["analyzed_at"] = datetime.now().isoformat()
    
    # 生成摘要
    if result.get("success") and result.get("content"):
        content = result["content"]
        result["summary"] = content[:200] + "..." if len(content) > 200 else content
        
        # 提取关键词（简单实现）
        words = content.split()
        result["keyword_count"] = len(words)
    
    # 保存到索引
    save_to_index(file_path, result)
    
    return result


def save_to_index(file_path: str, analysis_result: dict):
    """保存分析结果到索引"""
    index = load_json(DOC_INDEX_FILE)
    
    file_hash = hashlib.md5(file_path.encode()).hexdigest()
    
    index[file_hash] = {
        "file_path": file_path,
        "file_name": analysis_result.get("file_name", ""),
        "file_type": analysis_result.get("file_type", ""),
        "file_size": analysis_result.get("file_size", 0),
        "analyzed_at": analysis_result.get("analyzed_at", ""),
        "word_count": analysis_result.get("word_count", 0),
        "summary": analysis_result.get("summary", "")[:100]
    }
    
    save_json(DOC_INDEX_FILE, index)


def search_documents_index(query: str) -> List[dict]:
    """搜索文档索引"""
    index = load_json(DOC_INDEX_FILE)
    
    results = []
    for doc_id, doc_info in index.items():
        if query.lower() in doc_info.get("file_name", "").lower() or \
           query.lower() in doc_info.get("summary", "").lower():
            results.append({"id": doc_id, **doc_info})
    
    return results


def get_document_list() -> List[dict]:
    """获取已分析的文档列表"""
    index = load_json(DOC_INDEX_FILE)
    
    return [{"id": doc_id, **doc_info} for doc_id, doc_info in index.items()]


# ========== 文档对比 ==========

def compare_documents(file_path1: str, file_path2: str) -> dict:
    """对比两个文档"""
    doc1 = analyze_document(file_path1)
    doc2 = analyze_document(file_path2)
    
    if not doc1.get("success") or not doc2.get("success"):
        return {"success": False, "error": "无法解析文档"}
    
    content1 = doc1.get("content", "")
    content2 = doc2.get("content", "")
    
    # 简单对比
    common_words = set(content1.split()) & set(content2.split())
    
    return {
        "success": True,
        "doc1": {
            "name": doc1.get("file_name"),
            "word_count": doc1.get("word_count", 0)
        },
        "doc2": {
            "name": doc2.get("file_name"),
            "word_count": doc2.get("word_count", 0)
        },
        "comparison": {
            "common_words_count": len(common_words),
            "size_diff": doc2.get("word_count", 0) - doc1.get("word_count", 0),
            "common_words": list(common_words)[:20]
        }
    }


# ========== 工具定义 ==========
DOC_ANALYSIS_TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "analyze_document",
            "description": "分析文档内容，支持PDF、Word、Excel、CSV、TXT、Markdown、JSON等格式",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档文件路径"
                    }
                },
                "required": ["file_path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "list_documents",
            "description": "列出已分析的文档",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_documents",
            "description": "搜索已分析的文档",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "搜索关键词"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "compare_documents",
            "description": "对比两个文档的差异",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path1": {
                        "type": "string",
                        "description": "第一个文档路径"
                    },
                    "file_path2": {
                        "type": "string",
                        "description": "第二个文档路径"
                    }
                },
                "required": ["file_path1", "file_path2"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "summarize_document",
            "description": "生成文档摘要",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档文件路径"
                    },
                    "max_length": {
                        "type": "integer",
                        "description": "摘要最大长度，默认200字"
                    }
                },
                "required": ["file_path"]
            }
        }
    }
]


def execute_doc_analysis_tool(tool_name: str, arguments: dict) -> str:
    """执行文档分析工具"""
    try:
        if tool_name == "analyze_document":
            file_path = arguments.get("file_path", "")
            result = analyze_document(file_path)
            return json.dumps(result, ensure_ascii=False, indent=2)
        
        elif tool_name == "list_documents":
            docs = get_document_list()
            return json.dumps({
                "success": True,
                "documents": docs,
                "total": len(docs)
            }, ensure_ascii=False, indent=2)
        
        elif tool_name == "search_documents":
            query = arguments.get("query", "")
            results = search_documents_index(query)
            return json.dumps({
                "success": True,
                "results": results,
                "total": len(results)
            }, ensure_ascii=False, indent=2)
        
        elif tool_name == "compare_documents":
            file_path1 = arguments.get("file_path1", "")
            file_path2 = arguments.get("file_path2", "")
            result = compare_documents(file_path1, file_path2)
            return json.dumps(result, ensure_ascii=False, indent=2)
        
        elif tool_name == "summarize_document":
            file_path = arguments.get("file_path", "")
            max_length = arguments.get("max_length", 200)
            result = analyze_document(file_path)
            
            if result.get("success"):
                content = result.get("content", "")
                summary = content[:max_length] + "..." if len(content) > max_length else content
                return json.dumps({
                    "success": True,
                    "file_name": result.get("file_name"),
                    "summary": summary,
                    "word_count": result.get("word_count", 0)
                }, ensure_ascii=False, indent=2)
            else:
                return json.dumps(result, ensure_ascii=False)
        
        else:
            return json.dumps({"error": f"未知工具: {tool_name}"}, ensure_ascii=False)
    
    except Exception as e:
        return json.dumps({"error": f"工具执行失败: {str(e)}"}, ensure_ascii=False)


# ========== 获取支持的格式 ==========

def get_supported_formats() -> str:
    """获取支持的文档格式"""
    formats = {
        "文本文件": [".txt", ".md", ".json"],
        "办公文档": [".docx", ".doc", ".xlsx", ".xls", ".csv"],
        "PDF文件": [".pdf"],
    }
    
    return json.dumps(formats, ensure_ascii=False, indent=2)
